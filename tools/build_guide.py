"""Builds the Room Config Builder guide from its Markdown source.

The guide used to be a .docx somebody edited by hand and exported to PDF, which
meant the PDF and the app drifted apart quietly: nothing in the repo said the
document was stale, and nobody could diff it. The source is now
`documentation/Room_Config_Builder_Guide.md` — plain text, reviewable in a pull
request — and this script renders it to the two formats people actually open:

    python tools/build_guide.py

    documentation/Room_Config_Builder_Guide.pdf
    Room_Config_Builder_Guide.docx
    documentation/Beginners_Guide_to_Room_Config.pdf
    documentation/Beginners_Guide_to_Room_Config.docx

There are two guides — the operation guide, which answers everything, and the
beginner's guide, which answers the first week — and with no argument this
builds both, because a change to the app is as likely to have dated one as the
other. `python tools/build_guide.py beginners` builds just that one.

All of them are overwritten in place, so the file names people have bookmarked
keep working.

The Markdown it understands is deliberately small — the subset a manual needs:

    # / ## / ###     headings (a `#` starts a new page in the PDF)
    paragraphs       blank-line separated
    - item           bullets, one level of nesting with two spaces
    1. item          numbered lists
    | a | b |        tables, with a |---|---| separator row
    ```              fenced code blocks
    **bold**  *italic*  `code`   inline styles
    > note           a callout box

Anything else is passed through as text rather than guessed at.
"""

from __future__ import annotations

import html
import io
import os
import re
import sys

from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import (
    BaseDocTemplate,
    Frame,
    KeepTogether,
    PageBreak,
    PageTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
)
from reportlab.platypus.tableofcontents import TableOfContents

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DOCS = os.path.join(ROOT, "documentation")

# THE GUIDES THIS BUILDS, by the name you pass on the command line.
#
# There are two, and they are two documents rather than one long and one short:
# the operation guide answers everything, and the beginner's guide answers the
# things somebody hits in their first week, in the order they hit them. Both are
# Markdown in `documentation/`, so a change to the app and the change to its
# documentation can travel in the same commit and be read as a diff.
#
# The .docx of each lands where its readers already look for it. The operation
# guide's has been at the top of the repo since before it had a source; the
# beginner's has always been in `documentation/`. Moving either would break the
# links people have.
GUIDES = {
    "operation": {
        "source": os.path.join(DOCS, "Room_Config_Builder_Guide.md"),
        "pdf": os.path.join(DOCS, "Room_Config_Builder_Guide.pdf"),
        "docx": os.path.join(ROOT, "Room_Config_Builder_Guide.docx"),
    },
    "beginners": {
        "source": os.path.join(DOCS, "Beginners_Guide_to_Room_Config.md"),
        "pdf": os.path.join(DOCS, "Beginners_Guide_to_Room_Config.pdf"),
        "docx": os.path.join(DOCS, "Beginners_Guide_to_Room_Config.docx"),
    },
}

# Where the current build is writing. Set by build_one() before anything
# renders, so the two builders below did not have to grow a parameter each.
PDF_OUT = GUIDES["operation"]["pdf"]
DOCX_OUT = GUIDES["operation"]["docx"]

WINDOWS_FONTS = os.path.join(os.environ.get("WINDIR", r"C:\Windows"), "Fonts")

INK = colors.HexColor("#1B1F24")
MUTED = colors.HexColor("#5B6572")
ACCENT = colors.HexColor("#1F5FA8")
RULE = colors.HexColor("#C9D2DC")
PANEL = colors.HexColor("#F2F5F9")


# ---------------------------------------------------------------------------
#  MARKDOWN -> BLOCKS
# ---------------------------------------------------------------------------


def parse(md: str) -> list[dict]:
    """Turns the source into a flat list of blocks, in reading order."""
    blocks: list[dict] = []
    lines = md.replace("\r\n", "\n").split("\n")
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Fenced code.
        if stripped.startswith("```"):
            i += 1
            code: list[str] = []
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code.append(lines[i])
                i += 1
            i += 1
            blocks.append({"kind": "code", "lines": code})
            continue

        # Headings.
        heading = re.match(r"^(#{1,3})\s+(.*)$", stripped)
        if heading:
            blocks.append(
                {
                    "kind": "heading",
                    "level": len(heading.group(1)),
                    "text": heading.group(2).strip(),
                }
            )
            i += 1
            continue

        # Tables.
        if stripped.startswith("|") and i + 1 < len(lines) and re.match(
            r"^\s*\|[\s:\-|]+\|\s*$", lines[i + 1]
        ):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                row = lines[i].strip().strip("|")
                if not re.match(r"^[\s:\-|]+$", row):
                    # A cell may contain a literal pipe, written \| — which is
                    # not a nicety here: the rule syntax the guide documents
                    # uses one to separate alternatives.
                    cells = re.split(r"(?<!\\)\|", row)
                    rows.append([c.strip().replace(r"\|", "|") for c in cells])
                i += 1
            blocks.append({"kind": "table", "rows": rows})
            continue

        # Callout.
        if stripped.startswith(">"):
            note = []
            while i < len(lines) and lines[i].strip().startswith(">"):
                note.append(lines[i].strip().lstrip(">").strip())
                i += 1
            blocks.append({"kind": "note", "text": " ".join(note)})
            continue

        # Lists: bullets and numbers, with one level of nesting.
        if re.match(r"^\s*([-*]|\d+\.)\s+", line):
            items = []
            ordered = bool(re.match(r"^\s*\d+\.\s+", line))
            while i < len(lines) and re.match(r"^\s*([-*]|\d+\.)\s+", lines[i]):
                indent = len(lines[i]) - len(lines[i].lstrip())
                text = re.sub(r"^\s*([-*]|\d+\.)\s+", "", lines[i]).strip()
                # A wrapped continuation line belongs to the item above it.
                while (
                    i + 1 < len(lines)
                    and lines[i + 1].strip()
                    and not re.match(r"^\s*([-*]|\d+\.)\s+", lines[i + 1])
                    and not lines[i + 1].startswith("#")
                    and not lines[i + 1].strip().startswith("|")
                    and lines[i + 1].startswith(" ")
                ):
                    i += 1
                    text += " " + lines[i].strip()
                items.append({"text": text, "depth": 1 if indent >= 2 else 0})
                i += 1
            blocks.append({"kind": "list", "ordered": ordered, "items": items})
            continue

        # A paragraph runs until the next blank line.
        para = [stripped]
        i += 1
        while i < len(lines) and lines[i].strip() and not re.match(
            r"^(#{1,3}\s|\s*([-*]|\d+\.)\s|\||>|```)", lines[i]
        ):
            para.append(lines[i].strip())
            i += 1
        blocks.append({"kind": "para", "text": " ".join(para)})

    return blocks


def _stash_code(text: str) -> tuple[str, list[str]]:
    """Takes `code spans` out of the way before the emphasis rules run.

    A manual is full of keys with stars and underscores in them —
    `power1_outlet_*` most of all — and left in place a single one of those
    turns the rest of the sentence italic and unbalances the markup.
    """
    spans: list[str] = []

    def keep(match: re.Match) -> str:
        spans.append(match.group(1))
        return f"\x00{len(spans) - 1}\x00"

    return re.sub(r"`(.+?)`", keep, text.replace(r"\|", "|")), spans


def inline_to_rl(text: str) -> str:
    """Inline Markdown to reportlab's mini-HTML."""
    stashed, spans = _stash_code(text)
    out = html.escape(stashed)
    out = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", out)
    out = re.sub(r"(?<!\*)\*(?!\s)(.+?)(?<!\s)\*(?!\*)", r"<i>\1</i>", out)
    return re.sub(
        r"\x00(\d+)\x00",
        lambda m: '<font face="GuideMono" size="9.5" color="#134074">'
        f"{html.escape(spans[int(m.group(1))])}</font>",
        out,
    )


def inline_to_runs(text: str) -> list[tuple[str, str]]:
    """Inline Markdown to (text, style) runs, for the .docx writer."""
    runs: list[tuple[str, str]] = []
    stashed, spans = _stash_code(text)
    text = re.sub(r"\x00(\d+)\x00", lambda m: f"`{spans[int(m.group(1))]}`", stashed)
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|(?<!\*)\*(?!\s)[^`]+?(?<!\s)\*(?!\*))")
    for piece in pattern.split(text):
        if not piece:
            continue
        if piece.startswith("**") and piece.endswith("**"):
            runs.append((piece[2:-2], "b"))
        elif piece.startswith("`") and piece.endswith("`"):
            runs.append((piece[1:-1], "code"))
        elif piece.startswith("*") and piece.endswith("*"):
            runs.append((piece[1:-1], "i"))
        else:
            runs.append((piece, ""))
    return runs


# ---------------------------------------------------------------------------
#  PDF
# ---------------------------------------------------------------------------


def register_fonts() -> tuple[str, str]:
    """Real TrueType faces, so an em dash or an arrow is not a black box."""
    faces = [
        ("GuideBody", "calibri.ttf", "calibrib.ttf", "calibrii.ttf", "calibriz.ttf"),
        ("GuideBody", "arial.ttf", "arialbd.ttf", "ariali.ttf", "arialbi.ttf"),
    ]
    body = None
    for name, regular, bold, italic, bolditalic in faces:
        path = os.path.join(WINDOWS_FONTS, regular)
        if not os.path.exists(path):
            continue
        pdfmetrics.registerFont(TTFont(name, path))
        pdfmetrics.registerFont(TTFont(name + "-Bold", os.path.join(WINDOWS_FONTS, bold)))
        pdfmetrics.registerFont(
            TTFont(name + "-Italic", os.path.join(WINDOWS_FONTS, italic))
        )
        pdfmetrics.registerFont(
            TTFont(name + "-BoldItalic", os.path.join(WINDOWS_FONTS, bolditalic))
        )
        pdfmetrics.registerFontFamily(
            name,
            normal=name,
            bold=name + "-Bold",
            italic=name + "-Italic",
            boldItalic=name + "-BoldItalic",
        )
        body = name
        break
    if body is None:
        body = "Helvetica"

    mono_path = os.path.join(WINDOWS_FONTS, "consola.ttf")
    if os.path.exists(mono_path):
        pdfmetrics.registerFont(TTFont("GuideMono", mono_path))
        pdfmetrics.registerFont(
            TTFont("GuideMono-Bold", os.path.join(WINDOWS_FONTS, "consolab.ttf"))
        )
        pdfmetrics.registerFontFamily(
            "GuideMono", normal="GuideMono", bold="GuideMono-Bold"
        )
        mono = "GuideMono"
    else:
        mono = "Courier"
    return body, mono


class GuideDoc(BaseDocTemplate):
    """Two page templates: a bare title page, then the body with a footer."""

    def __init__(self, filename, title, **kw):
        super().__init__(filename, pagesize=LETTER, title=title, author="", **kw)
        frame = Frame(
            0.9 * inch,
            0.85 * inch,
            LETTER[0] - 1.8 * inch,
            LETTER[1] - 1.75 * inch,
            id="body",
        )
        self.addPageTemplates(
            [
                PageTemplate(id="title", frames=[frame]),
                PageTemplate(id="body", frames=[frame], onPage=self._chrome),
            ]
        )
        self.doc_title = title
        self.current_chapter = ""

    def _chrome(self, canvas, doc):
        canvas.saveState()
        canvas.setFont(BODY_FONT, 8)
        canvas.setFillColor(MUTED)
        canvas.drawString(0.9 * inch, 0.55 * inch, self.doc_title)
        canvas.drawRightString(
            LETTER[0] - 0.9 * inch, 0.55 * inch, f"Page {canvas.getPageNumber()}"
        )
        canvas.setStrokeColor(RULE)
        canvas.setLineWidth(0.5)
        canvas.line(
            0.9 * inch, 0.72 * inch, LETTER[0] - 0.9 * inch, 0.72 * inch
        )
        canvas.restoreState()

    def afterFlowable(self, flowable):
        """Feeds the table of contents from the headings as they are laid out."""
        if not isinstance(flowable, Paragraph):
            return
        style = flowable.style.name
        if style not in ("H1", "H2"):
            return
        level = 0 if style == "H1" else 1
        text = re.sub(r"<[^>]+>", "", flowable.getPlainText())
        key = f"toc-{id(flowable)}"
        self.canv.bookmarkPage(key)
        self.notify("TOCEntry", (level, text, self.page, key))


def build_pdf(blocks: list[dict], title: str, subtitle: str, tagline: str) -> None:
    ss = getSampleStyleSheet()
    body = ParagraphStyle(
        "Body",
        parent=ss["BodyText"],
        fontName=BODY_FONT,
        fontSize=10.5,
        leading=15.5,
        textColor=INK,
        spaceAfter=7,
        alignment=TA_LEFT,
    )
    styles = {
        "body": body,
        "h1": ParagraphStyle(
            "H1",
            parent=body,
            fontName=BODY_FONT + "-Bold" if BODY_FONT != "Helvetica" else "Helvetica-Bold",
            fontSize=20,
            leading=25,
            textColor=ACCENT,
            spaceBefore=2,
            spaceAfter=10,
        ),
        "h2": ParagraphStyle(
            "H2",
            parent=body,
            fontName=BODY_FONT + "-Bold" if BODY_FONT != "Helvetica" else "Helvetica-Bold",
            fontSize=14,
            leading=19,
            textColor=INK,
            spaceBefore=14,
            spaceAfter=5,
        ),
        "h3": ParagraphStyle(
            "H3",
            parent=body,
            fontName=BODY_FONT + "-Bold" if BODY_FONT != "Helvetica" else "Helvetica-Bold",
            fontSize=11.5,
            leading=16,
            textColor=colors.HexColor("#33404F"),
            spaceBefore=10,
            spaceAfter=3,
        ),
        "bullet": ParagraphStyle(
            "Bullet", parent=body, leftIndent=16, bulletIndent=4, spaceAfter=4
        ),
        "bullet2": ParagraphStyle(
            "Bullet2", parent=body, leftIndent=32, bulletIndent=20, spaceAfter=3
        ),
        "code": ParagraphStyle(
            "Code",
            parent=body,
            fontName=MONO_FONT,
            fontSize=9,
            leading=12.5,
            textColor=colors.HexColor("#14324F"),
            spaceAfter=0,
            spaceBefore=0,
        ),
        "cell": ParagraphStyle("Cell", parent=body, fontSize=9.5, leading=13, spaceAfter=0),
        "cellhead": ParagraphStyle(
            "CellHead",
            parent=body,
            fontName=BODY_FONT + "-Bold" if BODY_FONT != "Helvetica" else "Helvetica-Bold",
            fontSize=9.5,
            leading=13,
            spaceAfter=0,
            textColor=colors.white,
        ),
        "note": ParagraphStyle(
            "Note", parent=body, fontSize=10, leading=14.5, spaceAfter=0
        ),
    }

    story: list = []

    # --- title page ---
    story.append(Spacer(1, 2.2 * inch))
    story.append(
        Paragraph(
            title,
            ParagraphStyle(
                "Title",
                parent=styles["h1"],
                fontSize=30,
                leading=36,
                textColor=ACCENT,
            ),
        )
    )
    story.append(
        Paragraph(
            subtitle,
            ParagraphStyle(
                "Subtitle", parent=body, fontSize=14, leading=20, textColor=INK
            ),
        )
    )
    story.append(Spacer(1, 0.25 * inch))
    story.append(
        Paragraph(
            tagline,
            ParagraphStyle(
                "Tagline", parent=body, fontSize=11, leading=17, textColor=MUTED
            ),
        )
    )
    story.append(PageBreak())

    # --- contents ---
    toc = TableOfContents()
    toc.levelStyles = [
        ParagraphStyle(
            "TOC1",
            parent=body,
            fontName=BODY_FONT + "-Bold" if BODY_FONT != "Helvetica" else "Helvetica-Bold",
            fontSize=11,
            leading=17,
            spaceBefore=8,
        ),
        ParagraphStyle("TOC2", parent=body, fontSize=10, leading=14, leftIndent=18),
    ]
    # Its own style name, so afterFlowable does not put the contents page
    # into the contents.
    story.append(
        Paragraph(
            "What is in here",
            ParagraphStyle("H1Plain", parent=styles["h1"]),
        )
    )
    story.append(toc)
    story.append(PageBreak())

    first_chapter = True
    for block in blocks:
        kind = block["kind"]
        if kind == "heading":
            level = block["level"]
            if level == 1:
                if not first_chapter:
                    story.append(PageBreak())
                first_chapter = False
                story.append(Paragraph(inline_to_rl(block["text"]), styles["h1"]))
                story.append(
                    Table(
                        [[""]],
                        colWidths=[LETTER[0] - 1.8 * inch],
                        rowHeights=[2],
                        style=TableStyle(
                            [("BACKGROUND", (0, 0), (-1, -1), ACCENT)]
                        ),
                    )
                )
                story.append(Spacer(1, 10))
            else:
                story.append(
                    Paragraph(
                        inline_to_rl(block["text"]),
                        styles["h2" if level == 2 else "h3"],
                    )
                )
        elif kind == "para":
            story.append(Paragraph(inline_to_rl(block["text"]), body))
        elif kind == "list":
            for index, item in enumerate(block["items"], start=1):
                style = styles["bullet2"] if item["depth"] else styles["bullet"]
                marker = f"{index}." if block["ordered"] else "\u2022"
                story.append(
                    Paragraph(inline_to_rl(item["text"]), style, bulletText=marker)
                )
            story.append(Spacer(1, 4))
        elif kind == "code":
            rows = [
                [Paragraph(
                    html.escape(line).replace(" ", "&nbsp;") or "&nbsp;", styles["code"]
                )]
                for line in block["lines"]
            ]
            table = Table(rows, colWidths=[LETTER[0] - 1.8 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), PANEL),
                        ("BOX", (0, 0), (-1, -1), 0.5, RULE),
                        ("LEFTPADDING", (0, 0), (-1, -1), 8),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
                        ("TOPPADDING", (0, 0), (-1, -1), 1),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 1),
                    ]
                )
            )
            story.append(Spacer(1, 4))
            story.append(table)
            story.append(Spacer(1, 9))
        elif kind == "table":
            head, *rest = block["rows"]
            width = LETTER[0] - 1.8 * inch
            cols = len(head)
            # First column a little wider: it is the name of the thing.
            widths = [width * 0.32] + [width * 0.68 / (cols - 1)] * (cols - 1) if cols > 1 else [width]
            data = [[Paragraph(inline_to_rl(c), styles["cellhead"]) for c in head]]
            data += [
                [Paragraph(inline_to_rl(c), styles["cell"]) for c in row] for row in rest
            ]
            table = Table(data, colWidths=widths, repeatRows=1)
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, 0), ACCENT),
                        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, PANEL]),
                        ("GRID", (0, 0), (-1, -1), 0.4, RULE),
                        ("VALIGN", (0, 0), (-1, -1), "TOP"),
                        ("LEFTPADDING", (0, 0), (-1, -1), 6),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
                        ("TOPPADDING", (0, 0), (-1, -1), 5),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                    ]
                )
            )
            story.append(Spacer(1, 4))
            story.append(table)
            story.append(Spacer(1, 10))
        elif kind == "note":
            cell = Paragraph(inline_to_rl(block["text"]), styles["note"])
            table = Table([[cell]], colWidths=[LETTER[0] - 1.8 * inch])
            table.setStyle(
                TableStyle(
                    [
                        ("BACKGROUND", (0, 0), (-1, -1), colors.HexColor("#FFF8E6")),
                        ("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#E0C48A")),
                        ("LINEBEFORE", (0, 0), (0, -1), 3, colors.HexColor("#D6A233")),
                        ("LEFTPADDING", (0, 0), (-1, -1), 10),
                        ("RIGHTPADDING", (0, 0), (-1, -1), 10),
                        ("TOPPADDING", (0, 0), (-1, -1), 8),
                        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
                    ]
                )
            )
            story.append(Spacer(1, 4))
            story.append(KeepTogether(table))
            story.append(Spacer(1, 10))

    doc = GuideDoc(PDF_OUT, title)
    # Two passes: the first works out what page every heading landed on, the
    # second lays the contents page out with those numbers in it.
    doc.multiBuild(story, canvasmaker=_canvas_with_body_template())
    print(f"PDF  -> {PDF_OUT}")


def _canvas_with_body_template():
    """The title page has no footer; everything after it does."""
    from reportlab.pdfgen import canvas as canvas_module

    class GuideCanvas(canvas_module.Canvas):
        def showPage(self):
            super().showPage()

    return GuideCanvas


# ---------------------------------------------------------------------------
#  DOCX
# ---------------------------------------------------------------------------


def build_docx(blocks: list[dict], title: str, subtitle: str, tagline: str) -> None:
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt, RGBColor

    document = docx.Document()
    normal = document.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)

    heading = document.add_paragraph()
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x1F, 0x5F, 0xA8)
    sub = document.add_paragraph()
    sub_run = sub.add_run(subtitle)
    sub_run.font.size = Pt(14)
    tag = document.add_paragraph()
    tag_run = tag.add_run(tagline)
    tag_run.italic = True
    tag_run.font.color.rgb = RGBColor(0x5B, 0x65, 0x72)
    document.add_page_break()

    def write_runs(paragraph, text):
        for piece, kind in inline_to_runs(text):
            run = paragraph.add_run(piece)
            if kind == "b":
                run.bold = True
            elif kind == "i":
                run.italic = True
            elif kind == "code":
                run.font.name = "Consolas"
                run.font.size = Pt(10)

    for block in blocks:
        kind = block["kind"]
        if kind == "heading":
            paragraph = document.add_heading(level=block["level"])
            write_runs(paragraph, block["text"])
        elif kind == "para":
            write_runs(document.add_paragraph(), block["text"])
        elif kind == "list":
            for item in block["items"]:
                style = "List Number" if block["ordered"] else "List Bullet"
                if item["depth"]:
                    style += " 2"
                try:
                    paragraph = document.add_paragraph(style=style)
                except KeyError:
                    paragraph = document.add_paragraph()
                write_runs(paragraph, item["text"])
        elif kind == "code":
            for line in block["lines"]:
                paragraph = document.add_paragraph()
                paragraph.paragraph_format.space_after = Pt(0)
                run = paragraph.add_run(line or " ")
                run.font.name = "Consolas"
                run.font.size = Pt(9.5)
        elif kind == "note":
            paragraph = document.add_paragraph()
            run = paragraph.add_run("Note  ")
            run.bold = True
            write_runs(paragraph, block["text"])
        elif kind == "table":
            head, *rest = block["rows"]
            table = document.add_table(rows=1, cols=len(head))
            table.style = "Light Grid Accent 1"
            for cell, text in zip(table.rows[0].cells, head):
                cell.text = ""
                write_runs(cell.paragraphs[0], text)
            for row in rest:
                cells = table.add_row().cells
                for cell, text in zip(cells, row):
                    cell.text = ""
                    write_runs(cell.paragraphs[0], text)
            document.add_paragraph()

    document.save(DOCX_OUT)
    print(f"DOCX -> {DOCX_OUT}")


# ---------------------------------------------------------------------------


def build_one(name: str) -> int:
    """Renders one entry of [GUIDES] to its PDF and .docx."""
    global PDF_OUT, DOCX_OUT

    guide = GUIDES[name]
    source = guide["source"]
    if not os.path.exists(source):
        print(f"No guide source at {source}", file=sys.stderr)
        return 1
    PDF_OUT = guide["pdf"]
    DOCX_OUT = guide["docx"]

    md = io.open(source, encoding="utf-8").read()

    # The first three non-empty lines are the cover: title, subtitle, tagline.
    front = [ln.strip() for ln in md.split("\n") if ln.strip()][:3]
    title = front[0].lstrip("# ").strip()
    subtitle = front[1].strip()
    tagline = front[2].strip()
    rest = md.split("\n")
    # Everything after the cover block (marked by the first '---' line).
    if "---" in [ln.strip() for ln in rest]:
        cut = [ln.strip() for ln in rest].index("---")
        md = "\n".join(rest[cut + 1 :])

    blocks = parse(md)
    build_pdf(blocks, title, subtitle, tagline)
    try:
        build_docx(blocks, title, subtitle, tagline)
    except ImportError:
        print("python-docx is not installed — PDF only.", file=sys.stderr)
    return 0


def main() -> int:
    # No argument builds BOTH, because the commonest reason to run this is that
    # the app changed - and a change to the app is as likely to have dated the
    # beginner's guide as the operation one. Naming a guide builds only it.
    wanted = sys.argv[1:] or list(GUIDES)
    unknown = [w for w in wanted if w not in GUIDES]
    if unknown:
        print(
            f"Unknown guide(s): {', '.join(unknown)}. "
            f"Try: {', '.join(GUIDES)}",
            file=sys.stderr,
        )
        return 1
    for name in wanted:
        failed = build_one(name)
        if failed:
            return failed
    return 0


BODY_FONT, MONO_FONT = register_fonts()

if __name__ == "__main__":
    raise SystemExit(main())
